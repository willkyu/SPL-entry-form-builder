import re
from docx import Document
import sys
import requests
import pymysql
from random import choice
from time import sleep

"""根据需要选择filename，第一个是拖拽文件打开时使用，调试时推荐使用第二个"""
filename = sys.argv[1]
#filename='C:\\Users\\49637\\Desktop\\PY Codes\\dist\\联赛ID（21年选拔赛第一场）.docx'
document = Document(filename)

print('      苏州宝可梦月赛报名表生成器 Ver2.2')
print('---------------------------------------------')
print('   该软件用于生成各赛季苏州宝可梦月赛报名表')
print('        需配合邵老师提供的报名表使用')
print('          使用前请仔细阅读README')
print('       英文报名表需要连接神百进行翻译')
print('          请确认可以正常访问神百')
print('   若出现翻译失败，请务必联系制作者完善程序')
print('        由willkyu制作 禁止商业用途')
print('      使用中若有任何问题请与制作者联系')
print('            联系QQ：496373158')
print('')
print('')
#paragraphs用于存储所有段落
paragraphs = document.paragraphs

#input_info将表格前的一些问题录入

IDnumber = input('联赛ID：')
paragraphs[3].add_run(IDnumber)

def input_info(paragraph):
    temp = input(paragraph.text)
    paragraph.add_run(temp)



#team_list存储队伍信息
#team_list = [[] for i in range(6)]

#team_list存储队伍信息
team_list = [[] for i in range(6)]

special_dic={'Unseen Fist':'无形拳','Disguise':'画皮'}
nature_dic={'Hardy':'勤奋','Lonely':'怕寂寞','Brave':'勇敢','Adamant':'固执','Naughty':'頑皮','Bold':'大胆','Docile':'坦率','Impish':'淘气','Lax':'乐天','Relaxed':'悠闲','Modest':'內斂','Mild':'慢吞吞','Bashful':'害羞','Rash':'马虎','Quiet':'冷静','Calm':'溫和','Gentle':'温顺','Careful':'慎重','Quirky':'浮躁','Sassy':'自大','Timid':'胆小','Hasty':'急躁','Jolly':'爽朗','Naive':'天真','Serious':'认真'}
pattern = re.compile('"wgPageName":"(.*?)[（("]', re.S)
user_agent_list=['Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.5 (KHTML, like Gecko)','Mozilla/5.0 (Windows NT 6.1; WOW64; rv:11.0) Gecko/20100101 Firefox/11.0','Mozilla/4.0 (compatible; MSIE 8.0; Windows NT 6.1; WOW64; Trident/4.0; SLCC2; .NET CLR 2.0.50727; .NET CLR 3.5.30729; .NET CLR 3.0.30729; Media Center PC 6.0; .NET4.0C; .NET4.0E)','Opera/9.80 (Windows NT 6.1; WOW64; U; zh-cn) Presto/2.10.229 Version/11.62']
def translate(eng):
    try:
        url='https://wiki.52poke.com/zh-hans/'+eng
        kv = {'user-agent': choice(user_agent_list)}
        r = requests.get(url, headers=kv)
        r.raise_for_status()
        print(eng+"翻译成功")
        r.encoding = r.apparent_encoding
        res=re.search(pattern,r.text)
        return res.group(1)
    except:
        print(eng+"翻译失败")
        return eng

def get_team_eng(team):
    #names = re.findall('(\S*?)-* @', team)

#==============================================
    names = re.findall('\s*(.*?) @', team)
    for index, item in enumerate(names):
    #item.replace(' ', '_');
        names[index], _, _ = names[index].partition('-')
        names[index], _, _ = names[index].partition('(')
#===============================================

    #names = re.findall('\s*(.*?) @', team)
    #print(names)
    
    natures = re.findall('(\S*?) Nature',team)
    abilities = re.findall('Ability: (.*?)  \nLevel',team)
    items = re.findall('@ (.*?)  \nAbility',team)
    #print(team)
    pms = team.split("\n\n")  
    #print(pms)     
    for i in range(6):
        team_list[i].append(names[i])
        team_list[i].append(natures[i])
        team_list[i].append(abilities[i])
        team_list[i].append(items[i])
        #print(str(i)+'   '+pms[i])
        moves = re.findall('- (.*)\s*', pms[i])
        #moves = re.findall('- (.*?)   ',pms[i])
        #print(moves)
        #temp = re.match('^.*- (.*?)\s*$',pms[i]).group(1)
        #moves.append(temp)
        team_list[i].append(moves)

    #translate
    print('......翻译中......')
    for i1,j1 in enumerate(team_list):
        for i0 in range(4):
            if i0==1:
                team_list[i1][i0]=nature_dic[team_list[i1][i0]]
                continue
            if team_list[i1][i0] in special_dic:
                team_list[i1][i0]=special_dic[team_list[i1][i0]]
            team_list[i1][i0]=translate(team_list[i1][i0].replace(' ','_'))
            sleep(0.05)
        for i2,j2 in enumerate(team_list[i1][4]):
            team_list[i1][4][i2]=translate(team_list[i1][4][i2].replace(' ','_'))
            sleep(0.05)
    print('.....翻译完成.....')


def get_team_chs(team):
    names = re.findall('(\S*?) @', team)
    natures = re.findall('性格: (\S*?)  \n个体值', team)
    abilities = re.findall('特性: (\S*?)  \n等级', team)
    items = re.findall('@ (\S*?)  \n特性', team)
    pms = team.split("\n\n")
    for i in range(6):
        team_list[i].append(names[i])
        team_list[i].append(natures[i])
        team_list[i].append(abilities[i])
        team_list[i].append(items[i])

        moves = re.findall('- (\S*?)\s', pms[i])
        #temp = re.match('^.*- (\S*?)\s*$', pms[i]).group(1)
        #moves.append(temp)
        team_list[i].append(moves)


#print(len(paragraphs))
#将表格前的所有问题录入，根据情况修改参数，range左闭右开
for i in range(4, len(paragraphs)-2):
    input_info(paragraphs[i])

#输入队伍
#team = rawinput('以ps队伍格式输入队伍：')

print('以ps队伍格式输入队伍(以#键结束)：')
team = ''
s = input()
while s != '#':
    team = team + s + '\n'
    s = input()


while 1:
    language = input('输入的是英文队伍还是中文队伍？(英文输入E，中文输入C)：')
    if language in ['E','C']:
        break
    else:
        print('输入错误！')


if language=='E':
    get_team_eng(team)
else:
    get_team_chs(team)

print(team_list)
tables = document.tables
for i in range(24):
    for j in [1, 3]:
        if(i % 8 < 4):
            tables[0].cell(i, j).text = team_list[(i+j//2*24)//8][i % 8]
        else:
            if len(team_list[(i+j//2*24)//8][4]) >= (i % 8 - 4):
                tables[0].cell(i, j).text = team_list[(i+j//2*24)//8][4][i % 8 - 4]

filename=filename.replace('联赛ID',IDnumber)
#保存
document.save(filename)

print('   已生成报名表‘' + filename + '’')
input('             ...按任意键退出...')
